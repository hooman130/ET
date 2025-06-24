from docx import Document
from docx.shared import Inches
import os


def generate_scatter_report(plots_dir, output_path="scatter_report.docx"):
    """Generate a DOCX report with scatter plots arranged in a table.

    Parameters
    ----------
    plots_dir : str
        Directory containing subfolders for each farm. Each subfolder should
        include images named ``scatter_day1.png``, ``scatter_day2.png`` and
        ``scatter_day3.png``.
    output_path : str, optional
        Path to the output DOCX file. Defaults to ``scatter_report.docx``.
    """

    document = Document()
    document.add_heading("Forecast Scatter Plots", level=1)

    # Determine farms (subdirectories) sorted alphabetically
    farms = [d for d in os.listdir(plots_dir) if os.path.isdir(os.path.join(plots_dir, d))]
    farms.sort()

    table = document.add_table(rows=1 + len(farms), cols=4)
    table.style = "Light List"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Farm"
    hdr_cells[1].text = "Day +1"
    hdr_cells[2].text = "Day +2"
    hdr_cells[3].text = "Day +3"

    for row_idx, farm in enumerate(farms, start=1):
        row_cells = table.rows[row_idx].cells
        row_cells[0].text = farm.replace("_", " ")

        for day in range(1, 4):
            img_path = os.path.join(plots_dir, farm, f"scatter_day{day}.png")
            if os.path.isfile(img_path):
                paragraph = row_cells[day].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(2))
            else:
                row_cells[day].text = "N/A"

    document.save(output_path)
    return output_path


if __name__ == "__main__":
    PLOTS_DIR = "plots_test_ET_2014-2025"
    OUTPUT_FILE = "scatter_plots_report.docx"
    if not os.path.exists(PLOTS_DIR):
        raise FileNotFoundError(f"Directory '{PLOTS_DIR}' not found")
    generate_scatter_report(PLOTS_DIR, OUTPUT_FILE)
    print(f"Report saved to {OUTPUT_FILE}")
